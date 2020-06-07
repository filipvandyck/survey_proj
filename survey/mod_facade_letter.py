#!/usr/bin/env python3
#
#create facade letters from template and survey IFH export 
# 
# requires:
#	- pandas
#       - python-docx 


#--------------------------------------------------------------
#imports

import os
import pandas as pd
import re
from docx import Document
from datetime import datetime
#import mod_rename_convert as ren
#--------------------------------------------------------------

def docx_replace_regex(doc_obj, regex , replace):
 
    for p in doc_obj.paragraphs:
        if regex.search(p.text):
            inline = p.runs
            # Loop added to work with runs (strings with same style)
            for i in range(len(inline)):
                if regex.search(inline[i].text):
                    text = regex.sub(replace, inline[i].text)
                    inline[i].text = text
 
    for table in doc_obj.tables:
        for row in table.rows:
            for cell in row.cells:
                docx_replace_regex(cell, regex , replace)

#--------------------------------------------------------------


INPUTFOLDER  = os.path.dirname(os.path.realpath(__file__)) + '/documents/FFL/Template/' 
OUTPUTFOLDER = os.path.dirname(os.path.realpath(__file__)) + '/documents/FFL/Template/docx/'
RENAMEDFOLDER = os.path.dirname(os.path.realpath(__file__)) + '/documents/FFL/Renamed/'

INPUT_FACADE_LETTER    = 'Facade_letter_SK_NL.docx'
INPUT_FACADE_LETTER_FR = 'Facade_letter_SK_FR.docx'
INPUT_SURVEY           = 'SURVEY.csv'
INPUT_PROJECT          = 'PROJECTNUMBERS.xlsx' 


# extracts columns + removes quotes from survey + processes the template in outputfolder
def process_facade_letter(survey,facade_letter,outputfolder=OUTPUTFOLDER):
    print("Select Project")

    projectfile = INPUTFOLDER + INPUT_PROJECT

    project = pd.read_excel(projectfile, index_col=0,dtype={'GSM PXS': object}) 
    project.columns = project.columns.str.replace(' ','_')

    for row in project.index:
        print("[" + str(row)+"]\t" + project['projectnummer'][row])


    answer = input("\n :: ") 

    CONTACTTEL = project["GSM_PXS"][int(answer)]
    CONTACTPERSOON = project["Verantwoordelijke_PXS"][int(answer)]
    CONTACTEMAIL = project["Mail_PXs"][int(answer)]
    VANDAAG = datetime.now().strftime("%d-%m-%Y")

    print("Datum Vandaag voor brief? [" + VANDAAG + "]")

    print ("[Y] = Yes")
    print ("[DATUM] = ")
    answer = input("\n :: ")
    answer = answer.upper()
    
    if answer != 'Y':
        VANDAAG = answer

    print("Contact gegevens brief::")
    print(CONTACTTEL)
    print(CONTACTPERSOON)
    print(CONTACTEMAIL)
    print(VANDAAG)

    print('processing:(remove quotes+extra columns):\n' + survey)

    print('only facade?')

    print ("[Y] = Yes")
    print ("[ANY] = All")
    answer = input("\n :: ")
    answer = answer.upper()

    survey = pd.read_csv(survey, delimiter=';', encoding = "ISO-8859-1",dtype={'Zip': object})
    survey.columns = survey.columns.str.replace(' ','_')

    survey_obj = survey.select_dtypes(['object'])
    
    survey[survey_obj.columns] = survey_obj.apply(lambda x: x.str.lstrip("'"))
    
    if answer == 'Y':
        survey = survey[survey['Wall_Mount'] == 'Y']

    survey = survey.drop_duplicates(subset='LAM_MK', keep='first')
    

    for index, row in survey.iterrows():
        STREET = row['Street']
        NR =  str(row['Nr']) + row['Suffix']

        ADRES  = STREET[0:20] + NR
        LAMKEY = str(row['LAM_MK'])
        POSTCODE = str(row['Zip'])
        PLAATS = row['Municipality']

        documentname = "Doctyp_FFL-Lamkey_" + LAMKEY + "-name_" + ADRES 
        print(documentname)

        doc = Document(facade_letter)
        
        REGEX = re.compile(r"<Street>") 
        docx_replace_regex(doc,REGEX, STREET)
        

        

        docx_replace_regex(doc, re.compile(r"<nr>") , NR)




        docx_replace_regex(doc, re.compile(r"<Lamkey>") , LAMKEY)
        

        
        docx_replace_regex(doc, re.compile(r"<Vandaag>") , VANDAAG)
  
        docx_replace_regex(doc, re.compile(r"<Contacttelefoonnummer>") , CONTACTTEL)
        docx_replace_regex(doc, re.compile(r"<Contact>") , CONTACTPERSOON)
        
        docx_replace_regex(doc, re.compile(r"<Postcode>") , POSTCODE)
        docx_replace_regex(doc, re.compile(r"<Plaats>") , PLAATS)
        
        docx_replace_regex(doc, re.compile(r"<Contactemail>") , CONTACTEMAIL)
        
        doc.save(outputfolder + documentname + ".docx")




#process_facade_letter(INPUTFOLDER + INPUT_SURVEY, INPUTFOLDER + INPUT_FACADE_LETTER)

#ren.convert_doc_to_pdf_from_directory(OUTPUTFOLDER,RENAMEDFOLDER)


